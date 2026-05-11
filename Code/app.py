# Flask backend for the IDS prototype. Runs the Random Forest classifier
# and talks to a local Mistral 7B through Ollama for chat, uncertain packet analysis, and incident report generation.

from flask import Flask, render_template, request, jsonify, send_file
import joblib
import pandas as pd
import numpy as np
import os
import io
import ollama
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)


# --- Model + metadata (loaded once at startup) ---
model = joblib.load('../models/ids_model.joblib')
encoder = joblib.load('../models/label_encoder.joblib')
features_list = joblib.load('../models/feature_names.joblib')


# --- Confidence thresholds for the three tier workflow ---
# 0.89 came from the confidence sweep in the evaluation chapter.
THREAT_THRESHOLD = 0.89
UNCERTAIN_UPPER = 0.89
UNCERTAIN_LOWER = 0.50

# Global session state (single-user prototype).
packet_pointer = 0
last_detected_threat = "None"
scan_history = []
pending_review = None

# Pre-loaded so /run_scan doesn't hit disk every call.
# production_network_data.csv is large (~139MB) and not included in the repo.
# Falls back to viva_demo_data.csv so the app works out of the box.
_prod_path = '../data/processed/production_network_data.csv'
demo_df = pd.read_csv('../data/processed/viva_demo_data.csv')
uncertain_df = pd.read_csv('../data/processed/uncertain_demo_data.csv')
production_df = pd.read_csv(_prod_path) if os.path.exists(_prod_path) else demo_df

DATA_STREAMS = {
    "production": production_df,
    "demo": demo_df,
    "uncertain": uncertain_df,
}


# --- Attack info used by the UI and as grounding context for the LLM ---
ATTACK_INFO = {
    "DoS Hulk": {
        "name": "DoS Hulk",
        "category": "Denial of Service",
        "description": "A volumetric denial-of-service tool that floods a web server with a high volume of unique HTTP GET requests. It generates obfuscated URLs to bypass caching and overwhelm server resources.",
        "impact": "Server becomes unresponsive to legitimate users. CPU and memory usage spike to 100%.",
        "response": "Rate-limit incoming connections, block the source IP range, and enable web application firewall rules."
    },
    "DoS GoldenEye": {
        "name": "DoS GoldenEye",
        "category": "Denial of Service",
        "description": "An HTTP denial-of-service tool that uses KeepAlive connections with hidden cache-busting parameters. It holds connections open as long as possible to exhaust server connection pools.",
        "impact": "Server runs out of available connections. New users cannot connect.",
        "response": "Set connection timeouts, limit concurrent connections per IP, deploy a reverse proxy."
    },
    "DoS Slowhttptest": {
        "name": "DoS Slowhttptest",
        "category": "Denial of Service",
        "description": "A slow-rate HTTP attack that sends requests or headers very slowly, keeping connections alive indefinitely. The server waits for the request to complete, tying up resources.",
        "impact": "Server connection pool exhaustion without high bandwidth usage. Difficult to detect with traditional volume-based thresholds.",
        "response": "Configure minimum data rate thresholds, set header/body receive timeouts, use a reverse proxy with buffering."
    },
    "DoS slowloris": {
        "name": "DoS Slowloris",
        "category": "Denial of Service",
        "description": "Opens multiple connections to the target server and keeps them open by sending partial HTTP headers at regular intervals. The server keeps waiting for the headers to complete.",
        "impact": "All available connection slots consumed. Server appears online but refuses new connections.",
        "response": "Limit connections per source IP, reduce connection timeout values, use event-driven web servers like Nginx."
    },
    "FTP-Patator": {
        "name": "FTP-Patator",
        "category": "Brute Force",
        "description": "A brute-force attack against FTP (File Transfer Protocol) services. The attacker systematically tries username/password combinations to gain unauthorised access to file servers.",
        "impact": "If successful, the attacker gains read/write access to files on the server. May lead to data exfiltration.",
        "response": "Implement account lockout policies, use key-based authentication, restrict FTP access by IP, monitor for repeated failed logins."
    },
    "SSH-Patator": {
        "name": "SSH-Patator",
        "category": "Brute Force",
        "description": "A brute-force attack against SSH (Secure Shell) services. The attacker tries many password combinations to gain remote command-line access to a server.",
        "impact": "If successful, the attacker gains full shell access to the server. This is a critical compromise.",
        "response": "Disable password authentication (use SSH keys only), implement fail2ban, change the default SSH port, restrict access by IP."
    },
    "Web Attack - Brute Force": {
        "name": "Web Attack - Brute Force",
        "category": "Web Application Attack",
        "description": "Systematically tries login credentials against a web application's authentication page. Uses automated tools to test thousands of username/password combinations.",
        "impact": "Potential account compromise, especially for users with weak passwords.",
        "response": "Implement CAPTCHA, account lockout after failed attempts, multi-factor authentication, rate limiting on login endpoints."
    },
    "Web Attack - XSS": {
        "name": "Web Attack - XSS (Cross-Site Scripting)",
        "category": "Web Application Attack",
        "description": "Injects malicious JavaScript code into web pages viewed by other users. The script executes in the victim's browser, potentially stealing session cookies or redirecting to malicious sites.",
        "impact": "Session hijacking, credential theft, defacement, redirection to phishing pages.",
        "response": "Sanitise all user input, implement Content Security Policy headers, use HttpOnly cookie flags."
    },
    "Web Attack - SQL Injection": {
        "name": "Web Attack - SQL Injection",
        "category": "Web Application Attack",
        "description": "Inserts malicious SQL commands through input fields to manipulate the application's database. Can extract, modify, or delete data directly.",
        "impact": "Complete database compromise, data exfiltration, authentication bypass, potential server takeover.",
        "response": "Use parameterised queries/prepared statements, validate input types, apply least-privilege database permissions."
    },
    "Heartbleed": {
        "name": "Heartbleed (CVE-2014-0160)",
        "category": "Vulnerability Exploit",
        "description": "Exploits a flaw in the OpenSSL heartbeat extension. The attacker sends a malformed heartbeat request that causes the server to leak up to 64KB of memory contents per request.",
        "impact": "Leaks encryption keys, passwords, session tokens, and other sensitive data from server memory.",
        "response": "Patch OpenSSL immediately, revoke and reissue SSL certificates, reset all user passwords."
    },
    "Infiltration": {
        "name": "Infiltration",
        "category": "Advanced Persistent Threat",
        "description": "A multi-stage attack where the attacker first gains a foothold (often via phishing or exploit), then moves laterally through the network to access high-value targets.",
        "impact": "Long-term unauthorised access, data exfiltration, potential for further compromise of internal systems.",
        "response": "Network segmentation, endpoint detection and response (EDR), monitor for unusual lateral movement patterns."
    },
    "BENIGN": {
        "name": "Benign Traffic",
        "category": "Normal",
        "description": "Normal, legitimate network traffic with no malicious characteristics detected.",
        "impact": "No impact. This is expected network behaviour.",
        "response": "No action required."
    }
}


# --- Tool stubs for the chat route ---
# Return sentinel strings; the frontend reacts to them and runs the matching UI action.

def start_ids_scan(stream_type: str):
    return "SCAN_TRIGGERED"


def isolate_network_segment():
    # Simulation only - a real deployment would push rules to a firewall/SDN.
    target = last_detected_threat
    print(f"\n[SIMULATED FIREWALL]: Isolating {target} segment.")
    return (
        f"[SIMULATION] Firewall rules updated. Traffic matching '{target}' "
        f"pattern routed to quarantine VLAN. Segment isolated. Awaiting "
        f"operator review."
    )


def explain_attack(attack_type: str):
    return "EXPLAIN_TRIGGERED"


def show_statistics():
    return "STATS_TRIGGERED"


@app.route('/')
def home():
    return render_template('index.html')


@app.route('/chat', methods=['POST'])
def handle_chat():
    # Keyword router first, LLM fallback at the bottom - keeps common commands fast.
    global packet_pointer, last_detected_threat
    data = request.json
    user_input = data.get('message', '')
    selected_stream = data.get('stream', 'production')

    lower_input = user_input.lower().strip()

    if any(kw in lower_input for kw in ['scan', 'monitor', 'start', 'check network', 'analyse', 'analyze']):
        return jsonify({"message": "SCAN_TRIGGERED", "icon": "scan"})

    if any(kw in lower_input for kw in ['isolate', 'quarantine', 'block', 'contain', 'stop']):
        result = isolate_network_segment()
        return jsonify({"message": result, "icon": "shield"})

    if any(kw in lower_input for kw in ['stat', 'summary', 'report', 'how many', 'count']):
        return jsonify({
            "message": "STATS_RESULT",
            "icon": "stats",
            "stats": get_session_stats()
        })

    # Direct attack name lookup (strip spaces so "DoSHulk" still matches).
    for attack_name in ATTACK_INFO:
        if attack_name.lower() in lower_input or attack_name.lower().replace(' ', '') in lower_input.replace(' ', ''):
            return jsonify({
                "message": "EXPLAIN_RESULT",
                "icon": "info",
                "attack_info": ATTACK_INFO[attack_name]
            })

    # "What is X" style - partial match on name tokens > 3 chars so short filler words don't false match.
    if any(kw in lower_input for kw in ['what is', 'explain', 'tell me about', 'describe', 'what does']):
        for attack_name in ATTACK_INFO:
            name_parts = attack_name.lower().split()
            if any(part in lower_input for part in name_parts if len(part) > 3):
                return jsonify({
                    "message": "EXPLAIN_RESULT",
                    "icon": "info",
                    "attack_info": ATTACK_INFO[attack_name]
                })

    # Fall-through: let Mistral answer.
    try:
        response = ollama.chat(
            model='mistral',
            messages=[
                {
                    'role': 'system',
                    'content': (
                        f'You are a concise security assistant. Current threat: {last_detected_threat}. '
                        f'Answer security questions briefly. One or two sentences maximum.'
                    )
                },
                {'role': 'user', 'content': user_input}
            ],
        )
        return jsonify({"message": response.message.content, "icon": "ai"})
    except Exception as e:
        # Ollama is down or model isn't pulled - give a hint instead of 500.
        print(f"[LLM ERROR]: {e}")
        return jsonify({
            "message": "Command not recognised. Try: 'scan network', 'isolate', 'what is DoS Hulk?', or 'show stats'.",
            "icon": "ai"
        })


@app.route('/run_scan', methods=['POST'])
def run_scan():
    data = request.json
    selected_stream = data.get('stream', 'production')
    return perform_inference(selected_stream)


@app.route('/get_stats', methods=['GET'])
def get_stats():
    return jsonify(get_session_stats())


@app.route('/get_attack_info/<attack_type>', methods=['GET'])
def get_attack_info(attack_type):
    for key in ATTACK_INFO:
        if key.lower() == attack_type.lower() or attack_type.lower() in key.lower():
            return jsonify(ATTACK_INFO[key])
    return jsonify({"error": f"Unknown attack type: {attack_type}"}), 404


@app.route('/analyse_packet', methods=['POST'])
def analyse_packet():
    # LLM analysis for the packet currently in the uncertain band.
    global pending_review
    if not pending_review:
        return jsonify({"analysis": "No packet pending review."})

    pkt = pending_review
    attack_info = ATTACK_INFO.get(pkt['label'], {})

    top3_str = ", ".join(
        f"{t['label']} ({t['confidence']}%)" for t in pkt.get('top3', [])
    )
    net = pkt.get('network_info', {})

    # Give Mistral the numeric context so it reasons from facts, not guesses.
    prompt = (
        f"A network intrusion detection system flagged a packet as potentially malicious, "
        f"but the confidence is only {pkt['confidence']}% - not high enough to auto-alert.\n\n"
        f"Classification: {pkt['label']} ({pkt['confidence']}% confidence)\n"
        f"Top-3 predictions: {top3_str}\n"
        f"Destination port: {net.get('dest_port', 'unknown')}\n"
        f"Flow duration: {net.get('flow_duration', 'unknown')}\n"
        f"Forward packets: {net.get('fwd_packets', 'unknown')}\n"
        f"Backward packets: {net.get('bwd_packets', 'unknown')}\n"
        f"Attack category: {pkt.get('category', 'Unknown')}\n\n"
        f"Known info about {pkt['label']}: {attack_info.get('description', 'No info available.')}\n\n"
        f"As a security analyst, explain in plain English:\n"
        f"1. Why the system flagged this packet\n"
        f"2. What the confidence split between the top predictions suggests\n"
        f"3. What an operator should look for to decide if this is a real threat or a false positive\n"
        f"Keep it concise - 3-4 sentences maximum."
    )

    try:
        response = ollama.chat(
            model='mistral',
            messages=[
                {
                    'role': 'system',
                    'content': 'You are a senior security analyst helping a junior operator review uncertain IDS alerts. Be clear, specific, and actionable.'
                },
                {'role': 'user', 'content': prompt}
            ],
        )
        analysis = response.message.content
    except Exception as e:
        # Rule based fallback so the operator still gets a usable answer.
        print(f"[LLM ANALYSIS ERROR]: {e}")
        analysis = (
            f"The system detected patterns consistent with {pkt['label']} "
            f"({pkt.get('category', 'unknown category')}), but at only {pkt['confidence']}% confidence. "
            f"The top prediction split is: {top3_str}. "
            f"A close split between threat and benign suggests this could be a false positive. "
            f"Check the destination port ({net.get('dest_port', '?')}) and packet count "
            f"({net.get('fwd_packets', '?')} fwd / {net.get('bwd_packets', '?')} bwd) "
            f"against normal baseline traffic for this segment."
        )

    return jsonify({
        "analysis": analysis,
        "packet_details": pkt
    })


@app.route('/resolve_packet', methods=['POST'])
def resolve_packet():
    # Operator decision on the pending uncertain packet.
    global pending_review, last_detected_threat, scan_history
    data = request.json
    decision = data.get('decision', '')

    if not pending_review:
        return jsonify({"message": "No packet pending review.", "icon": "info"})

    pkt = pending_review
    packet_num = pkt['packet']

    # Flip the history entry so the session stats reflect the resolution.
    for entry in scan_history:
        if entry['packet'] == packet_num:
            if decision == 'threat':
                entry['is_threat'] = True
                entry['is_uncertain'] = False
                entry['operator_resolved'] = 'threat'
                last_detected_threat = pkt['label']
            else:
                entry['is_threat'] = False
                entry['is_uncertain'] = False
                entry['operator_resolved'] = 'dismissed'
            break

    pending_review = None

    if decision == 'threat':
        return jsonify({
            "message": f"Packet #{packet_num} marked as THREAT ({pkt['label']}). You can now isolate this segment.",
            "icon": "alert",
            "resolved_as": "threat"
        })
    else:
        return jsonify({
            "message": f"Packet #{packet_num} dismissed as false positive.",
            "icon": "safe",
            "resolved_as": "dismissed"
        })


@app.route('/generate_report', methods=['POST'])
def generate_report():
    # LLM writes the narrative; fallback produces a plain-text version if Ollama is down.
    if not scan_history:
        return jsonify({"report": "No scan data available. Run a scan first."})

    stats = get_session_stats()

    threats = [s for s in scan_history if s.get('is_threat')]
    uncertain = [s for s in scan_history if s.get('is_uncertain')]
    resolved = [s for s in scan_history if s.get('operator_resolved')]

    # Cap each list at 10 so the prompt stays inside Mistral's context.
    events_summary = ""
    for t in threats[:10]:
        events_summary += (
            f"- Packet #{t['packet']}: {t['label']} detected at "
            f"{t['confidence']*100:.1f}% confidence (Port {t.get('dest_port', '?')})\n"
        )

    for u in uncertain[:10]:
        resolution = u.get('operator_resolved', 'pending')
        events_summary += (
            f"- Packet #{u['packet']}: {u['label']} at "
            f"{u['confidence']*100:.1f}% confidence - UNCERTAIN, "
            f"operator resolved as: {resolution}\n"
        )

    prompt = f"""Write a formal security incident report based on the following IDS scan session data.

SESSION SUMMARY:
- Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}
- Total packets scanned: {stats['total_scanned']}
- Confirmed threats: {stats['threats_detected']}
- Uncertain detections reviewed: {stats['uncertain_count']}
- Benign packets: {stats['benign_count']}
- Average threat confidence: {stats['avg_confidence']}%
- Threat types detected: {', '.join(f"{k} ({v})" for k, v in stats['threat_breakdown'].items()) if stats['threat_breakdown'] else 'None'}

NOTABLE EVENTS:
{events_summary if events_summary else 'No notable events.'}

OPERATOR DECISIONS:
{len(resolved)} uncertain packets were manually reviewed.
{sum(1 for r in resolved if r.get('operator_resolved') == 'threat')} escalated to confirmed threat.
{sum(1 for r in resolved if r.get('operator_resolved') == 'dismissed')} dismissed as false positive.

Write the report with these sections:
1. EXECUTIVE SUMMARY (2-3 sentences overview)
2. INCIDENT TIMELINE (chronological list of key detections)
3. THREAT ANALYSIS (what attack types were found, their severity, and what they mean)
4. OPERATOR ACTIONS (what the operator reviewed and decided)
5. RISK ASSESSMENT (overall risk level: Critical/High/Medium/Low, with justification)
6. RECOMMENDATIONS (specific next steps)

Use formal, professional language suitable for management and compliance documentation. Be specific with numbers and percentages."""

    try:
        response = ollama.chat(
            model='mistral',
            messages=[
                {
                    'role': 'system',
                    'content': 'You are a senior security analyst writing a formal incident report for management and compliance records. Be precise, professional, and actionable.'
                },
                {'role': 'user', 'content': prompt}
            ],
        )
        report_text = response.message.content
    except Exception as e:
        print(f"[LLM REPORT ERROR]: {e}")
        report_text = generate_fallback_report(stats, threats, uncertain, resolved)

    return jsonify({
        "report": report_text,
        "stats": stats
    })


@app.route('/export_report', methods=['POST'])
def export_report():
    # Render the report text into a .docx and send it back.
    data = request.json
    report_text = data.get('report', '')

    if not report_text:
        return jsonify({"error": "No report content provided."}), 400

    stats = get_session_stats()
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    title = doc.add_heading('Security Incident Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    doc.add_paragraph('')
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Light Shading Accent 1'
    meta_data = [
        ('Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ('Generated By', 'Specialised Security Partner - IDS Prototype'),
        ('Classification Engine', 'Random Forest (scikit-learn) on CICIDS2017'),
        ('Analysis Engine', 'Mistral 7B (Local LLM via Ollama)'),
        ('Data Sovereignty', 'All processing performed locally - no data sent to external services'),
    ]
    for i, (key, val) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = key
        meta_table.rows[i].cells[1].text = val
        for cell in meta_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph('')

    doc.add_heading('Session Statistics', level=2)
    stats_table = doc.add_table(rows=1, cols=5)
    stats_table.style = 'Medium Shading 1 Accent 1'
    headers = ['Packets Scanned', 'Threats', 'Uncertain', 'Benign', 'Avg Confidence']
    for i, h in enumerate(headers):
        stats_table.rows[0].cells[i].text = h
        for p in stats_table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    row = stats_table.add_row()
    values = [
        str(stats['total_scanned']),
        str(stats['threats_detected']),
        str(stats['uncertain_count']),
        str(stats['benign_count']),
        f"{stats['avg_confidence']}%",
    ]
    for i, v in enumerate(values):
        row.cells[i].text = v
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

    doc.add_paragraph('')

    doc.add_heading('Analysis Report', level=2)

    lines = report_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        is_heading = False
        for prefix in ['1.', '2.', '3.', '4.', '5.', '6.']:
            if line.startswith(prefix):
                heading_text = line[2:].strip().strip(':').strip('*').strip('#')
                doc.add_heading(heading_text, level=3)
                is_heading = True
                break

        if not is_heading:
            if line.startswith('#'):
                heading_text = line.lstrip('#').strip().strip(':').strip('*')
                doc.add_heading(heading_text, level=3)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
            else:
                clean = line.replace('**', '').replace('__', '')
                doc.add_paragraph(clean)

    doc.add_paragraph('')
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Generated by Specialised Security Partner - All data processed locally')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"incident_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


def generate_fallback_report(stats, threats, uncertain, resolved):
    report = f"""SECURITY INCIDENT REPORT
Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}

1. EXECUTIVE SUMMARY
A network scan session processed {stats['total_scanned']} packets, detecting {stats['threats_detected']} confirmed threats and {stats['uncertain_count']} uncertain detections requiring operator review. {len(resolved)} packets were manually reviewed by the operator.

2. INCIDENT TIMELINE
"""
    for t in threats[:10]:
        report += (
            f"- Packet #{t['packet']}: {t['label']} detected "
            f"({t['confidence']*100:.1f}% confidence, Port {t.get('dest_port', '?')})\n"
        )
    for u in uncertain[:10]:
        res = u.get('operator_resolved', 'pending')
        report += (
            f"- Packet #{u['packet']}: {u['label']} "
            f"({u['confidence']*100:.1f}% confidence) - Operator: {res}\n"
        )

    report += f"""
3. THREAT ANALYSIS
Threat types detected: {', '.join(f"{k} ({v} occurrences)" for k, v in stats['threat_breakdown'].items()) if stats['threat_breakdown'] else 'None'}
Average threat confidence: {stats['avg_confidence']}%

4. OPERATOR ACTIONS
{sum(1 for r in resolved if r.get('operator_resolved') == 'threat')} uncertain packets escalated to confirmed threats.
{sum(1 for r in resolved if r.get('operator_resolved') == 'dismissed')} uncertain packets dismissed as false positives.

5. RISK ASSESSMENT
Overall risk level: {'CRITICAL' if stats['threats_detected'] > 10 else 'HIGH' if stats['threats_detected'] > 5 else 'MEDIUM' if stats['threats_detected'] > 0 else 'LOW'}

6. RECOMMENDATIONS
- Review all confirmed threat sources and update firewall rules accordingly.
- Investigate uncertain detections that were escalated for further forensic analysis.
- Consider updating IDS signatures if false positive rate exceeds acceptable threshold.
"""
    return report


@app.route('/reset', methods=['POST'])
def reset_scan():
    global packet_pointer, last_detected_threat, scan_history, pending_review
    packet_pointer = 0
    last_detected_threat = "None"
    scan_history = []
    pending_review = None
    return jsonify({"message": "Session reset. Ready for new scan."})


def get_session_stats():
    if not scan_history:
        return {
            "total_scanned": 0,
            "threats_detected": 0,
            "uncertain_count": 0,
            "benign_count": 0,
            "threat_breakdown": {},
            "avg_confidence": 0,
            "highest_confidence": 0,
        }

    threats = [s for s in scan_history if s['is_threat']]
    uncertain = [s for s in scan_history if s.get('is_uncertain', False)]
    benign = [s for s in scan_history if not s['is_threat'] and not s.get('is_uncertain', False)]

    threat_types = {}
    confidences = []
    for t in threats:
        label = t['label']
        threat_types[label] = threat_types.get(label, 0) + 1
        confidences.append(t['confidence'])

    return {
        "total_scanned": len(scan_history),
        "threats_detected": len(threats),
        "uncertain_count": len(uncertain),
        "benign_count": len(benign),
        "threat_breakdown": threat_types,
        "avg_confidence": round(np.mean(confidences) * 100, 1) if confidences else 0,
        "highest_confidence": round(max(confidences) * 100, 1) if confidences else 0,
    }


def perform_inference(stream_key):
    # Classify the next packet and route it through the three-tier workflow.
    global packet_pointer, last_detected_threat, pending_review

    try:
        df = DATA_STREAMS.get(stream_key)
        if df is None:
            return jsonify({"message": "Invalid stream selected.", "icon": "error"})

        if packet_pointer >= len(df):
            packet_pointer = 0
            return jsonify({
                "message": "End of data stream reached. Pointer reset.",
                "icon": "reset",
                "stats": get_session_stats()
            })

        row = df.iloc[[packet_pointer]]
        packet_pointer += 1

        input_data = row[features_list]
        probs = model.predict_proba(input_data)[0]
        max_prob = np.max(probs)
        pred_idx = np.argmax(probs)
        pred_label = str(encoder.inverse_transform([pred_idx])[0])

        top3_idx = np.argsort(probs)[-3:][::-1]
        top3 = []
        for idx in top3_idx:
            label = str(encoder.inverse_transform([idx])[0])
            prob = float(probs[idx])
            if prob > 0.01:
                top3.append({"label": label, "confidence": round(prob * 100, 1)})

        dest_port = int(row['Destination Port'].values[0]) if 'Destination Port' in row.columns else 0
        flow_duration = float(row['Flow Duration'].values[0]) if 'Flow Duration' in row.columns else 0
        fwd_packets = int(row['Total Fwd Packets'].values[0]) if 'Total Fwd Packets' in row.columns else 0
        bwd_packets = int(row['Total Backward Packets'].values[0]) if 'Total Backward Packets' in row.columns else 0

        is_threat = max_prob >= THREAT_THRESHOLD and pred_label.upper() != "BENIGN"
        is_uncertain = (
            not is_threat
            and max_prob >= UNCERTAIN_LOWER
            and max_prob < UNCERTAIN_UPPER
            and pred_label.upper() != "BENIGN"
        )

        network_info = {
            "dest_port": dest_port,
            "flow_duration": flow_duration,
            "fwd_packets": fwd_packets,
            "bwd_packets": bwd_packets,
        }

        scan_history.append({
            'packet': packet_pointer,
            'label': pred_label,
            'confidence': max_prob,
            'is_threat': is_threat,
            'is_uncertain': is_uncertain,
            'dest_port': dest_port,
        })

        if is_threat:
            last_detected_threat = pred_label
            return jsonify({
                "message": f"ALERT: {pred_label} detected",
                "icon": "alert",
                "packet": packet_pointer,
                "confidence": round(max_prob * 100, 1),
                "label": pred_label,
                "is_threat": True,
                "is_uncertain": False,
                "top3": top3,
                "network_info": network_info,
                "category": ATTACK_INFO.get(pred_label, {}).get("category", "Unknown"),
            })

        if is_uncertain:
            pending_review = {
                "packet": packet_pointer,
                "label": pred_label,
                "confidence": round(max_prob * 100, 1),
                "top3": top3,
                "network_info": network_info,
                "category": ATTACK_INFO.get(pred_label, {}).get("category", "Unknown"),
            }
            return jsonify({
                "message": f"UNCERTAIN: {pred_label} - {round(max_prob * 100, 1)}% confidence. Operator review required.",
                "icon": "uncertain",
                "packet": packet_pointer,
                "confidence": round(max_prob * 100, 1),
                "label": pred_label,
                "is_threat": False,
                "is_uncertain": True,
                "top3": top3,
                "network_info": network_info,
                "category": ATTACK_INFO.get(pred_label, {}).get("category", "Unknown"),
            })

        return jsonify({
            "message": "Benign",
            "icon": "safe",
            "packet": packet_pointer,
            "confidence": round(max_prob * 100, 1),
            "label": pred_label,
            "is_threat": False,
            "is_uncertain": False,
            "top3": top3,
            "network_info": network_info,
        })

    except Exception as e:
        return jsonify({"message": f"Inference Error: {str(e)}", "icon": "error"})


if __name__ == '__main__':
    app.run(debug=True, port=5000)
